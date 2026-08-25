from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from dotenv import load_dotenv
import os
from docx import Document as DocxDocument

load_dotenv()

# --- LOAD YOUR DOCX ---
file_path = "data/policy.docx"

print(f"Loading DOCX: {file_path}")
docx = DocxDocument(file_path)
full_text = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
docs = [Document(page_content=full_text)]

print(f"Loaded {len(full_text)} characters")

splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
chunks = splitter.split_documents(docs)
print(f"Split into {len(chunks)} chunks")

# Use GOOGLE FREE embeddings (no torch needed)
embeddings = GoogleGenerativeAIEmbeddings(
    model="gemini-embedding-001",
    google_api_key=os.getenv("GOOGLE_API_KEY")
)

print("Creating Chroma DB... this takes 20 sec...")
db = Chroma.from_documents(chunks, embeddings, persist_directory="./chroma_db")

print("DONE! DB saved in ./chroma_db")